from flask import Flask, request, render_template, send_file
import pandas as pd
from openai import OpenAI
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# --- OpenAI Embedding function ---
def get_embedding(text, model="text-embedding-3-small"):
    try:
        response = client.embeddings.create(input=[text], model=model)
        return response.data[0].embedding
    except Exception as e:
        print("Embedding error:", e)
        return None

# --- Static company details ---
COMPANY_DETAILS = {
    "COMPANY NAME": "Travel Lykke Private Limited",
    "MAIL ID": "support@lykke.travel",
    "PHONE NO.": "+91-8047493335",
    "ADDRESS": "1st Floor, Bhau Institute – COEP, Shivajinagar, Pune, Maharashtra 411005"
}

# --- Egypt escalation + notes ---
EGYPT_ESCALATION = """\
Escalation Matrix India Office:
Leader (Ms. Heba): +201025267265
Sales Leader (Lalit): +918826894140

Escalation Matrix Egypt Office:
Egypt Operations Head (Ms. Reem Soud): +2 01011900567
Company CEO (Mr. Hossam): +2 01116101030
"""

EGYPT_NOTES = """\
Special Note on Egypt:
* Total Payable tip for the trip: 120 USD [Days in Egypt*5*Number of Pax], kindly pay the same upon arrival.
* Tips are cultural and compulsory in Egypt. If you had to pay the same extra anywhere beyond the amount collected in advance, please intimate the team and the same will be refunded.
* Check-in/Checkout timing, hotel bedding type, etc. are controlled by respective accommodation only.
"""

EGYPT_COMMENTS = """\
Other Comments or Special Instructions:
Above mentioned services are confirmed and non-changeable.
Any changes would attract full cancellation of tours/services and additional tours/services will be charged as per the company policies and directly payable in Egypt only.
"""

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        try:
            # --- Uploads + Inputs ---
            master_file = request.files["master_file"]
            itinerary_text = request.form["itinerary"].strip()
            confirmation_id = request.form.get("confirmation_id", "N/A")
            passenger_name = request.form.get("passenger_name", "N/A")
            start_date = request.form.get("start_date", "N/A")
            destination = request.form.get("destination", "N/A")

            df_master = pd.read_excel(master_file)
            master_services = df_master["Service"].tolist()
            master_outputs = df_master["Formatted Output"].tolist()
            master_hotels = df_master.get("Hotel", pd.Series([""] * len(df_master))).tolist()

            # Precompute embeddings
            master_embeddings = [get_embedding(s) for s in master_services]

            # Parse itinerary lines
            itinerary_lines = itinerary_text.split("\n")
            document = Document()

            # Insert Logo
            try:
                document.add_picture("static/logo.png", width=Inches(2))
            except Exception:
                document.add_paragraph("[Logo Missing]")

            # --- Company Details Section ---
            document.add_heading("COMPANY DETAILS", level=1)
            company_table = document.add_table(rows=len(COMPANY_DETAILS), cols=2)
            company_table.style = "Table Grid"
            for i, (k, v) in enumerate(COMPANY_DETAILS.items()):
                company_table.cell(i, 0).text = k
                company_table.cell(i, 1).text = v
            document.add_paragraph("")

            # --- Booking Details Section ---
            document.add_heading("BOOKING DETAILS", level=1)
            booking_table = document.add_table(rows=4, cols=2)
            booking_table.style = "Table Grid"
            booking_table.cell(0, 0).text = "Confirmation ID"
            booking_table.cell(0, 1).text = confirmation_id
            booking_table.cell(1, 0).text = "Passenger / Guest Name"
            booking_table.cell(1, 1).text = passenger_name
            booking_table.cell(2, 0).text = "Travel Dates"
            booking_table.cell(2, 1).text = f"{start_date} to TBD"
            booking_table.cell(3, 0).text = "Destination"
            booking_table.cell(3, 1).text = destination
            document.add_paragraph("")

            # --- Day-wise itinerary ---
            document.add_heading("ITINERARY", level=1)
            for i, line in enumerate(itinerary_lines, 1):
                document.add_heading(f"Day {i}", level=2)
                services = [s.strip() for s in line.split("+")]

                for service in services:
                    query_emb = get_embedding(service)
                    if query_emb:
                        sims = [sum(a*b for a, b in zip(query_emb, ref)) for ref in master_embeddings]
                        best_idx = sims.index(max(sims))
                        formatted = master_outputs[best_idx]
                        hotel = master_hotels[best_idx]
                    else:
                        formatted = f"⚠️ Could not process: {service}"
                        hotel = ""

                    document.add_paragraph(f"Service: {formatted}")
                    if hotel:
                        document.add_paragraph(f"Hotel: {hotel}")

            # --- Escalation + Notes + Comments ---
            document.add_heading("Escalation Matrix", level=1)
            document.add_paragraph(EGYPT_ESCALATION)

            document.add_heading("Special Notes", level=1)
            document.add_paragraph(EGYPT_NOTES)

            document.add_heading("Other Comments", level=1)
            document.add_paragraph(EGYPT_COMMENTS)

            # --- Save file ---
            output = BytesIO()
            document.save(output)
            output.seek(0)
            return send_file(output, as_attachment=True, download_name="voucher.docx")

        except Exception as e:
            return f"❌ Error: {str(e)}"

    return render_template("index.html")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
